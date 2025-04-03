from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from typing import List, Optional
import os
import shutil
from datetime import datetime
import markdown
import docx
import requests
from bs4 import BeautifulSoup

from app.database import get_db
from app.models.article import Article, ArticleImage, ArticleCategory, ArticleStatus

router = APIRouter()

# Article CRUD operations
@router.post("/articles/")
async def create_article(
    title: str,
    content: Optional[str] = None,
    category_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    article = Article(
        title=title,
        content=content,
        category_id=category_id,
        status=ArticleStatus.DRAFT
    )
    db.add(article)
    db.commit()
    db.refresh(article)
    return article

@router.get("/articles/")
async def get_articles(
    skip: int = 0,
    limit: int = 10,
    status: Optional[str] = None,
    category_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    query = db.query(Article)
    if status:
        query = query.filter(Article.status == status)
    if category_id:
        query = query.filter(Article.category_id == category_id)
    
    total = query.count()
    articles = query.offset(skip).limit(limit).all()
    
    return {
        "total": total,
        "items": articles
    }

@router.get("/articles/{article_id}")
async def get_article(article_id: int, db: Session = Depends(get_db)):
    article = db.query(Article).filter(Article.id == article_id).first()
    if not article:
        raise HTTPException(status_code=404, detail="Article not found")
    return article

@router.put("/articles/{article_id}")
async def update_article(
    article_id: int,
    title: Optional[str] = None,
    content: Optional[str] = None,
    category_id: Optional[int] = None,
    status: Optional[str] = None,
    db: Session = Depends(get_db)
):
    article = db.query(Article).filter(Article.id == article_id).first()
    if not article:
        raise HTTPException(status_code=404, detail="Article not found")
    
    if title:
        article.title = title
    if content:
        article.content = content
        # Convert markdown to HTML
        article.content_html = markdown.markdown(content, extensions=['fenced_code', 'tables', 'mermaid'])
    if category_id:
        article.category_id = category_id
    if status:
        article.status = status
    
    article.updated_at = datetime.now()
    db.commit()
    db.refresh(article)
    return article

@router.delete("/articles/{article_id}")
async def delete_article(article_id: int, db: Session = Depends(get_db)):
    article = db.query(Article).filter(Article.id == article_id).first()
    if not article:
        raise HTTPException(status_code=404, detail="Article not found")
    
    article.status = ArticleStatus.DELETED
    db.commit()
    return {"message": "Article deleted successfully"}

# File upload and import
@router.post("/articles/upload/image")
async def upload_image(
    file: UploadFile = File(...),
    article_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    # Create uploads directory if it doesn't exist
    upload_dir = "uploads/images"
    os.makedirs(upload_dir, exist_ok=True)
    
    # Save file
    file_path = f"{upload_dir}/{datetime.now().strftime('%Y%m%d%H%M%S')}_{file.filename}"
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Save to database if article_id is provided
    if article_id:
        image = ArticleImage(
            article_id=article_id,
            url=file_path
        )
        db.add(image)
        db.commit()
        db.refresh(image)
        return image
    
    return {"url": file_path}

@router.post("/articles/import/word")
async def import_word(
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    # Save uploaded file temporarily
    temp_file = f"temp_{file.filename}"
    with open(temp_file, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Read Word document
    doc = docx.Document(temp_file)
    title = doc.paragraphs[0].text
    content = "\n".join([p.text for p in doc.paragraphs[1:]])
    
    # Create article
    article = Article(
        title=title,
        content=content,
        source_type="word",
        content_html=markdown.markdown(content)
    )
    db.add(article)
    db.commit()
    db.refresh(article)
    
    # Clean up
    os.remove(temp_file)
    
    return article

@router.post("/articles/import/url")
async def import_from_url(
    url: str,
    db: Session = Depends(get_db)
):
    # Fetch content from URL
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')
    
    # Extract title and content
    title = soup.title.string if soup.title else "Imported Article"
    content = soup.get_text()
    
    # Create article
    article = Article(
        title=title,
        content=content,
        source_type="url",
        source_url=url,
        content_html=content
    )
    db.add(article)
    db.commit()
    db.refresh(article)
    
    return article

# Category operations
@router.post("/categories/")
async def create_category(
    name: str,
    parent_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    category = ArticleCategory(name=name, parent_id=parent_id)
    db.add(category)
    db.commit()
    db.refresh(category)
    return category

@router.get("/categories/")
async def get_categories(db: Session = Depends(get_db)):
    return db.query(ArticleCategory).all()

@router.put("/categories/{category_id}")
async def update_category(
    category_id: int,
    name: str,
    parent_id: Optional[int] = None,
    db: Session = Depends(get_db)
):
    category = db.query(ArticleCategory).filter(ArticleCategory.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Category not found")
    
    category.name = name
    if parent_id is not None:
        category.parent_id = parent_id
    
    db.commit()
    db.refresh(category)
    return category

@router.delete("/categories/{category_id}")
async def delete_category(category_id: int, db: Session = Depends(get_db)):
    category = db.query(ArticleCategory).filter(ArticleCategory.id == category_id).first()
    if not category:
        raise HTTPException(status_code=404, detail="Category not found")
    
    # Check if category has articles
    articles_count = db.query(Article).filter(Article.category_id == category_id).count()
    if articles_count > 0:
        raise HTTPException(status_code=400, detail="Cannot delete category with articles")
    
    db.delete(category)
    db.commit()
    return {"message": "Category deleted successfully"} 